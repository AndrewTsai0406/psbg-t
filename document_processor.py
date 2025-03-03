import re
import ast
import glob
import pandas as pd
import zipfile
from tqdm import tqdm
from pathlib import Path
from lxml import etree
from typing import Tuple
from docx import Document
from docx.table import Table
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches

from rag_pipleline import RAG


from prompts import (
    style_transfer_prompt_question_mark,
    style_transfer_prompt_question_mark_v2,
    table_style_transfer_prompt,
)

from customed_block_func_map import (
    customed_block_func_map,
    hash_with_fixed_seed,
    DOE_COC_calculation,
)
from langchain_core.utils.json import parse_json_markdown
import fuzzy_json
from typing import List


class DocumentProcessor:
    """
    Class to process a document and perform style transfer on its content.
    """

    def __init__(self, pair, **kwargs):
        """
        Initialize the DocumentProcessor.

        Args:
            pair (tuple): A tuple containing the target and source document paths.
        """
        self.pair = pair
        self.doc = Document(pair[1])
        # Counter for each heading level
        self.level_counters = {}
        # Stores the visible number of each paragraph
        self.list_items = {}
        # Indicates the start of a new block
        self.start_flag = False
        # Ancestors of the current block
        self.ancestors = ""
        # Stores Para that belong to the same block
        self.temporary_context = [[]]
        # Para for RAG: [element_ind:int, "":str, self.ancestors:str]
        self.paragraphs_to_rag = []
        # Tables for RAG: [element_ind:int, df:pd.DataFrame, table_columns_width:list[int], html_table:str, self.ancestors:str]
        self.tables_to_rag = []
        # Modified Para to add: [self.paragraphs_to_rag[i][0]:int, response:str]
        self.paragraphs_to_add = []
        # Para to remove: [element_ind:int, para_ind:int]
        self.paragraphs_to_remove = []
        # Modified Tables to add: [element_ind:int, responsed_table:pd.DataFrame, table_columns_width:list[int]]
        self.tables_to_add = []
        # Tables to remove: [element_ind:int, table_ind:int]
        self.tables_to_remove = []
        # Initialize a RAG pipeline
        self.RAG: RAG = None
        self.kwargs = kwargs

    def process_document(self):
        """Process the document by handling its elements and performing style transfer."""
        self.process_elements()
        self.perform_style_transfer_paragraphs()
        self.perform_style_transfer_tables()
        self.perform_modifications()
        self.handle_footer(
            model_no=self.kwargs.get("model_no", "Default Model No"),
            date=self.kwargs.get("date", "Default Date"),
            drawn=self.kwargs.get("drawn", "Default Drawn"),
            design_ee=self.kwargs.get("design_ee", "Default Design (EE)"),
            design_me=self.kwargs.get("design_me", "Default Design (ME)"),
            document_name=self.kwargs.get("document_name", "Default Document Name"),
            rev=self.kwargs.get("rev", "Default Rev"),
        )
        self.save_modified_document()

    def process_elements(self):
        """Iterate over the document elements and process paragraphs and tables."""
        for element_ind, child in enumerate(self.doc.element.body):
            if (
                self.pair[1] in customed_block_func_map
                and hash_with_fixed_seed(str(child.xml))
                in customed_block_func_map[self.pair[1]].keys()
            ):
                function_name = customed_block_func_map[self.pair[1]][
                    hash_with_fixed_seed(str(child.xml))
                ]
                print(f"Using customed function {function_name}")
                globals()[function_name](self, element_ind, child)
            else:
                if isinstance(child, CT_P):
                    self.process_paragraph(element_ind, child)
                elif isinstance(child, CT_Tbl):
                    self.process_table(element_ind, child)

    def process_paragraph(self, element_ind, child):
        """
        Process a paragraph element.
        """
        for para_ind, paragraph in enumerate(self.doc.paragraphs):
            if paragraph._element == child and paragraph.text:
                visible_number, item_or_text, level = self.get_visible_number(
                    paragraph, self.pair[1], self.level_counters
                )
                if not visible_number or "•" in visible_number:
                    self.paragraphs_to_remove.append([element_ind, para_ind])
                    if self.start_flag:
                        self.temporary_context[-1].append(paragraph.text)
                else:
                    self.start_flag = True
                    self.list_items[visible_number] = re.sub(
                        r"^[0-9]+(?:\.[0-9]+)*\.?\s*", "", paragraph.text
                    )
                    if self.paragraphs_to_rag:
                        self.paragraphs_to_rag[-1][1] += "\n".join(
                            self.temporary_context[-1]
                        )
                    self.ancestors = self.get_ancestors(self.list_items, visible_number)
                    self.paragraphs_to_rag.append([element_ind, "", self.ancestors])
                    self.temporary_context.append([])
                break

    def process_table(self, element_ind, child):
        """
        Process a table element.
        """
        for table_ind, table in enumerate(self.doc.tables):
            if table._element == child:
                self.tables_to_remove.append([element_ind, table_ind])
                df, table_columns_width, html_table = self.extract_table_info(table)
                if table_ind:
                    self.tables_to_rag.append(
                        [
                            element_ind,
                            df,
                            table_columns_width,
                            html_table,
                            self.ancestors,
                        ]
                    )
                break

    def extract_table_info(self, table):
        """
        Extract data from a table.
        """
        table_rows = []
        table_columns_width = []
        for row in table.rows:
            if not table_columns_width:
                table_columns_width = [cell.width for cell in row.cells]
            row_data = [cell.text for cell in row.cells]
            table_rows.append(row_data)
        df = pd.DataFrame(table_rows[1:], columns=table_rows[0])
        grid = self.extract_table_data(table)
        html_table = self.generate_html_table(grid)
        return df, table_columns_width, html_table

    def perform_style_transfer_paragraphs(self):
        """Apply style transfer to the paragraphs to be added."""
        ancestors_lst = [item[2] for item in self.paragraphs_to_rag]
        templates = [item[1] for item in self.paragraphs_to_rag]
        # templates = [re.sub(r"\d", "?", template) for template in templates]
        responses = []
        for template, ancestors in tqdm(
            zip(templates, ancestors_lst),
            total=len(templates),
            desc="Text Content RAG Inference",
        ):
            if len(template) > 1 and bool(re.search(r"\d", template)):  # 有數字才詢問
                response = self.RAG.ask_rag(
                    template,
                    ancestors,
                    mode="text",
                    use_constrain=False,
                )
                try:
                    response = parse_json_markdown(response, parser=fuzzy_json.loads)
                    response = response["revised_template"]
                    responses.append(response)
                except:
                    responses.append(template)
            else:
                responses.append(template)

        for i, response in enumerate(responses):
            self.paragraphs_to_add.append([self.paragraphs_to_rag[i][0], response])

    from pandas import DataFrame

    def perform_style_transfer_tables(self):
        """Apply style transfer to the tables to be added."""
        for ind, (
            element_ind,
            ori_df,  # type: DataFrame
            table_columns_width,
            html_table,
            ancestors,
        ) in tqdm(
            enumerate(self.tables_to_rag),
            total=len(self.tables_to_rag),
            desc="Table RAG Inference",
        ):
            # df, ori_cols = self.rename_duplicated_columns(ori_df)
            cols = pd.Series(ori_df.columns)
            ori_cols = list(cols.copy())
            ori_df = ori_df.astype(str)

            # if not len(df) or len(df.to_string()) >= 6000:
            #     print("Using original table due to size constraints.")
            #     responsed_table = ori_df.replace(r"\d", "?", regex=True)
            # else:
            #     responsed_table = self.attempt_style_transfer_on_table(
            #         df, ori_cols, ancestors
            #     )
            #     if responsed_table is None:
            #         responsed_table = ori_df.replace(r"\d", "?", regex=True)

            responsed_table = self.attempt_style_transfer_on_table(
                ori_df, ori_cols, ancestors
            )
            if responsed_table is None:
                responsed_table = ori_df.replace(r"\d", "?", regex=True)

            responsed_table.reset_index(drop=True, inplace=True)
            responsed_table.columns = ori_cols
            self.tables_to_add.append(
                [
                    element_ind,
                    responsed_table,
                    table_columns_width,
                ]
            )

    def attempt_style_transfer_on_table(
        self, ori_df: pd.DataFrame, ori_cols, ancestors
    ):
        """
        Attempt to apply style transfer to a table.
        """
        for _ in range(3):
            try:

                response = self.RAG.ask_rag(ori_df, ancestors, mode="table")
                table_dict = ast.literal_eval(
                    response[response.index("{") : response.rfind("}") + 1]
                )
                responsed_table = pd.DataFrame(table_dict)
                responsed_table.columns = ori_cols
                return responsed_table
            except Exception as e:
                print(f"Single-shot style transfer failed: {e}, retrying...")

        return None

    def perform_modifications(self):
        """
        Apply the collected modifications to the document.
        """
        para_to_reserve = set()
        for lst_add in self.paragraphs_to_add:
            if "?" in lst_add[1]:
                para_to_reserve.add(lst_add[0] + 1)
        self.paragraphs_to_remove = [
            remove_lst
            for remove_lst in self.paragraphs_to_remove
            if not remove_lst[0] in para_to_reserve
        ]

        table_to_reserve = set()
        table_to_reserve.add(0)
        for lst_add in self.tables_to_add:
            if lst_add[1].map(lambda x: "?" in str(x)).any().any():
                table_to_reserve.add(lst_add[0])
        self.tables_to_remove = [
            remove_lst
            for remove_lst in self.tables_to_remove
            if not remove_lst[0] in table_to_reserve
        ]

        content_to_add = [p + ["paragraph"] for p in self.paragraphs_to_add] + [
            t + ["table"] for t in self.tables_to_add
        ]
        content_to_remove = [p + ["paragraph"] for p in self.paragraphs_to_remove] + [
            t + ["table"] for t in self.tables_to_remove
        ]

        self.handle_modification(self.doc, content_to_add, content_to_remove)

    def handle_footer(self, **kwargs):
        """
        Handle footer replacements.
        """
        item_locatoryKey_pair = {
            "model_no": "MODEL NO",
            "date": "Date",
            "drawn": "Drawn",
            "design_ee": "Design (EE)",
            "design_me": "Design (ME)",
            "document_name": "DOCUMENT NAME",
            "rev": "REV",
        }

        def find_and_replace_footer_text(table: Table, text: str) -> bool:
            for row_idx in range(len(table.rows) - 1, -1, -1):
                row = table.rows[row_idx]
                for col_idx in range(len(row.cells) - 1, -1, -1):
                    cell = row.cells[col_idx]

                    if locatoryKey in cell.text:
                        if len(cell.paragraphs) > 1:  # model_no or document_name
                            row_index = row_idx
                            column_index = col_idx
                            para_index = -1
                        elif len(cell.paragraphs) == 1:
                            row_index = row_idx + 1
                            column_index = col_idx
                            para_index = 0

                        p = (
                            table.rows[row_index]
                            .cells[column_index]
                            .paragraphs[para_index]
                            ._p
                        )
                        for run in (
                            table.rows[row_index]
                            .cells[column_index]
                            .paragraphs[para_index]
                            .runs[:-1]
                        ):
                            p.remove(run._r)

                        table.rows[row_index].cells[column_index].paragraphs[
                            para_index
                        ].runs[0].text = text

                        return True
            return False

        footer_table = self.doc.sections[0].footer.tables[0]
        first_table = self.doc.tables[0]
        first_table_have_footer = True

        for item, locatoryKey in item_locatoryKey_pair.items():
            replace_text = kwargs.get(item)
            find_and_replace_footer_text(footer_table, replace_text)
            if first_table_have_footer:
                first_table_have_footer = find_and_replace_footer_text(
                    first_table, replace_text
                )

    def save_modified_document(self):
        """Save the modified document to a new file."""
        old_path = Path(self.pair[1])

        new_root = Path("./modified_data")
        new_path = new_root / old_path.relative_to(old_path.parts[0]).parent
        new_file = new_path / (old_path.stem + "_modified" + old_path.suffix)
        new_file.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(new_file)

    # -----------------------------
    # Internal Utility Methods
    # -----------------------------

    def rename_duplicated_columns(self, df):
        cols = pd.Series(df.columns)
        ori_cols = list(cols.copy())
        for dup in cols[cols.duplicated()].unique():
            cols[cols[cols == dup].index.values.tolist()] = [
                dup + "_" + chr(96 + i) if i != 0 else dup
                for i in range(sum(cols == dup))
            ]
        df.columns = cols
        return df, ori_cols

    def get_ancestors(self, list_items, number) -> List[str]:
        ancestors = []
        for i in range(len(number.split(".")), 0, -1):
            ancestor = ".".join(number.split(".")[:i])
            if ancestor in list_items:
                content = list_items[ancestor]
                if not "." in ancestor:
                    ancestor += "."
                ancestors.append([ancestor, content])

        return [" ".join(anscestor) for anscestor in ancestors[::-1]]

    def get_numbering_definitions(self, docx_file: str) -> etree.Element:
        with zipfile.ZipFile(docx_file) as docx:
            with docx.open("word/numbering.xml") as numbering_file:
                xml_content = numbering_file.read()
                return etree.fromstring(xml_content)

    def get_visible_number(
        self, paragraph, file_path, level_counters
    ) -> Tuple[str, str, str]:
        """
        Returns a tuple containing:
        1) The visible numbering string if the paragraph is part of a list (e.g., "1.1"),
        2) The type of numbering ("item" if recognized as a numbered/bulleted list, "text" if inferred from paragraph text),
        3) The nesting level of the list.
        If no valid numbering is detected, returns empty strings.

        Explanation of main points:

        • get_numbering_definitions(file_path): Gathers the Word XML structure that defines how paragraphs are numbered or bulleted in the given docx file.
        • paragraph._p.pPr.numPr: Accesses the paragraph's numbering properties in the underlying lxml/_Element interface of python-docx.
        • level_counters: A dictionary that keeps track of the current count at each outline level (e.g., the second bullet under the first item might be level_counters[1] = 2, assuming level 0 is the top level).
        • text_list_item(paragraph.text): This method presumably checks whether the paragraph string starts with something like "1.", "2)", etc., to detect a list pattern if the docx numbering properties aren’t set.
        • Returns either:
        – A formatted string of digits for numbered lists (e.g., "1.2.1"),
        – A bullet point ("•") if it’s recognized as a bulleted list, or
        – Nothing at all if the paragraph isn’t recognized as a list item.
        """

        # Retrieve XML representation of the numbering definitions from the docx file
        numbering_definitions = self.get_numbering_definitions(file_path)

        # Define the XML namespaces so the XPath queries can reference the w: prefix
        NAMESPACES = {
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        }

        # Check if the paragraph has paragraph properties (pPr) and a defined numbering property (numPr)
        if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
            # Get the numbering ID and level from the paragraph's XML
            num_id = paragraph._p.pPr.numPr.numId.val
            level = paragraph._p.pPr.numPr.ilvl.val

            # Find the matching abstractNumId for this numId in the numbering definitions
            num_def = numbering_definitions.xpath(
                f"//w:num[@w:numId='{num_id}']/w:abstractNumId", namespaces=NAMESPACES
            )

            # If a match is found, retrieve its abstractNumId value
            if num_def:
                abstract_num_id = num_def[0].get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                )

                # Get the <w:numFmt> node that indicates the numbering format (e.g., 'decimal', 'bullet', etc.)
                level_def = numbering_definitions.xpath(
                    f"//w:abstractNum[@w:abstractNumId='{abstract_num_id}']/w:lvl[@w:ilvl='{level}']/w:numFmt",
                    namespaces=NAMESPACES,
                )

                # If we found a numbering format, determine what it is
                if level_def:
                    num_format = level_def[0].get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                    )

                    # If the numbering format is decimal, increment the corresponding level counter
                    if num_format == "decimal":
                        # If this level is not yet in the dictionary, default to 0 before incrementing
                        level_counters[level] = level_counters.get(level, 0) + 1

                        # Reset counters for deeper levels when a higher level (lower number) increments
                        for k in level_counters:
                            if k > level:
                                level_counters[k] = 0

                    # Build the partial or full numbering string for all levels (e.g., "1.2.1")
                    visible_number = []
                    for lvl in range(level + 1):
                        # Check if the level counter exists for this level
                        if lvl in level_counters:
                            # For decimal numbering, append the numeric value; for bullet, use a bullet point
                            if num_format == "decimal":
                                visible_number.append(str(level_counters[lvl]))
                            else:
                                visible_number.append(
                                    "•"
                                )  # A simple bullet placeholder

                    # If any numbering was assembled, return the joined string, the type, and the current level
                    if visible_number:
                        return (
                            ".".join(visible_number),
                            "item",
                            level,
                        )

        # If the paragraph isn't recognized by docx as a list, check if the text itself may be a list item (like "1. " or "a. ")
        text_numbering = self.text_list_item(paragraph.text)
        if text_numbering:
            # If so, return that text as the numbering, "text" type, and the nesting level based on the count of dots
            return text_numbering, "text", text_numbering.count(".")

        # If none of the above conditions are satisfied, return empty values
        return "", "", ""

    def extract_table_data(self, table):
        grid = []
        merged_cells = {}

        for row_idx, row in enumerate(table.rows):
            grid_row = []
            for col_idx, cell in enumerate(row.cells):
                cell_key = (row_idx, col_idx)
                if cell_key in merged_cells:
                    continue

                cell_text = cell.text.strip()
                rowspan, colspan = self.get_span(cell)

                for i in range(rowspan):
                    for j in range(colspan):
                        if not (i == 0 and j == 0):
                            merged_cells[(row_idx + i, col_idx + j)] = True

                cell_data = {"text": cell_text, "rowspan": rowspan, "colspan": colspan}
                grid_row.append(cell_data)
            grid.append(grid_row)

        return grid

    def get_span(self, cell):
        tc = cell._tc
        grid_span = tc.xpath(".//w:gridSpan")
        v_merge = tc.xpath(".//w:vMerge")

        colspan = int(grid_span[0].get(qn("w:val"))) if grid_span else 1
        rowspan = 1

        if v_merge:
            v_merge_val = v_merge[0].get(qn("w:val"))
            if v_merge_val == "restart":
                rowspan = 1
                next_tc = tc.getnext()
                while next_tc is not None:
                    next_v_merge = next_tc.xpath(".//w:vMerge")
                    if next_v_merge and next_v_merge[0].get(qn("w:val")) is None:
                        rowspan += 1
                        next_tc = next_tc.getnext()
                    else:
                        break
            elif v_merge_val is None:
                rowspan = 0

        return rowspan, colspan

    def generate_html_table(self, grid):
        html = '<table border="1">\n'
        for row in grid:
            html += "  <tr>\n"
            for cell in row:
                rowspan = f' rowspan="{cell["rowspan"]}"' if cell["rowspan"] > 1 else ""
                colspan = f' colspan="{cell["colspan"]}"' if cell["colspan"] > 1 else ""
                html += f'    <td{rowspan}{colspan}>{cell["text"]}</td>\n'
            html += "  </tr>\n"
        html += "</table>"
        return html

    def text_list_item(self, text):
        text = text.strip()
        if text.endswith("."):
            text = text + "0"
        pattern_0 = r"^\d+\.(\d+\.)*\d+\.*\s+"
        pattern_1 = r"^\d+\.+\s+"
        combined_pattern = f"({pattern_0}|{pattern_1})"
        result = re.findall(combined_pattern, text)
        return max(result[0], key=len).strip() if result else None

    def handle_modification(self, doc, content_to_add, content_to_remove):
        content_to_add = sorted(content_to_add, reverse=True, key=lambda x: x[0])
        content_to_remove = sorted(content_to_remove, reverse=True, key=lambda x: x[0])

        for idx_remove, _ in enumerate(content_to_remove):
            if content_to_remove[idx_remove][-1] == "paragraph":
                self.remove_paragraph(content_to_remove[idx_remove][1], doc)
            elif content_to_remove[idx_remove][-1] == "table":
                self.remove_table(content_to_remove[idx_remove][1], doc)

            for idx_add, _ in enumerate(content_to_add):
                if content_to_add[idx_add][0] >= content_to_remove[idx_remove][0]:
                    content_to_add[idx_add][0] -= 1

        for each_add in content_to_add:
            if each_add[-1] == "paragraph":
                self.add_paragraph_text(each_add[0] + 1, doc, f"{each_add[1]}")
            elif each_add[-1] == "table":
                self.add_table(each_add[0] + 1, doc, each_add[1], each_add[2])

    def remove_paragraph(self, i, doc):
        par = doc.paragraphs[i]._element
        par.getparent().remove(par)
        par._p = par._element = None

    def remove_table(self, i, doc):
        tables = doc.tables
        tbl = tables[i]
        tbl._element.getparent().remove(tbl._element)

    def add_paragraph_text(self, i, doc, new_text):
        doc_segs = re.split(r"(\?+)", new_text)
        new_para = doc.add_paragraph()
        new_para.paragraph_format.left_indent = Inches(1)
        unsure_text = any([True for seg in doc_segs if "?" in seg])

        for seg in doc_segs:
            new_run = new_para.add_run(seg)
            if "?" in seg:
                highlight = OxmlElement("w:highlight")
                highlight.set(qn("w:val"), "yellow")
                new_run._r.get_or_add_rPr().append(highlight)
            else:
                if unsure_text:
                    highlight = OxmlElement("w:highlight")
                    highlight.set(qn("w:val"), "lightGray")
                    new_run._r.get_or_add_rPr().append(highlight)

        doc._body._element.insert(i, new_para._element)

    def add_table(self, i, doc, df, table_columns_width):
        table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))

        # Add columns
        for ind, column in enumerate(df.columns):
            table.cell(0, ind).text = column

        # Add rows
        for indx, (_, row) in enumerate(df.iterrows()):
            for indy, (_, cell_value) in enumerate(row.items()):
                try:
                    table.cell(indx + 1, indy).text = str(cell_value)
                except:
                    table.cell(indx + 1, indy).text = "N/A"
                for paragraph in table.cell(indx + 1, indy).paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        self.set_column_widths(table, table_columns_width)
        self.set_table_indent(table)
        self.set_cell_border(table)
        self.merge_same_cells(table)
        self.center_table_cells(table)
        self.color_table_cells(table)

        doc._body._element.insert(i, table._element)

    def center_table_cells(self, table):
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def set_table_indent(self, table, indent=1440):
        tbl_pr = table._element.tblPr
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        if tbl_ind is None:
            tbl_ind = OxmlElement("w:tblInd")
            tbl_pr.append(tbl_ind)
        tbl_ind.set(qn("w:w"), str(indent))
        tbl_ind.set(qn("w:type"), "dxa")

    def set_cell_border(self, table):
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._element.tcPr
                if tc_pr is None:
                    tc_pr = OxmlElement("w:tcPr")
                    cell._element.append(tc_pr)

                tc_borders = OxmlElement("w:tcBorders")

                for border_name in ["top", "left", "bottom", "right"]:
                    border = OxmlElement(f"w:{border_name}")
                    border.set(qn("w:val"), "single")
                    border.set(qn("w:sz"), "4")
                    border.set(qn("w:space"), "0")
                    border.set(qn("w:color"), "000000")
                    tc_borders.append(border)

                tc_pr.append(tc_borders)

    def merge_same_cells(self, table):
        # Horizontal merging
        for row in table.rows:
            cells = row.cells
            num_cells = len(cells)
            i = 0
            while i < num_cells:
                cell_text = cells[i].text.strip()
                start_idx = i
                i += 1
                while i < num_cells and cells[i].text.strip() == cell_text:
                    i += 1
                end_idx = i - 1
                if end_idx > start_idx:
                    start_cell = cells[start_idx]
                    for merge_idx in range(start_idx + 1, end_idx + 1):
                        try:
                            start_cell.merge(cells[merge_idx])
                            cells[merge_idx].text = ""
                        except:
                            pass
                    start_cell.text = cell_text

    def color_table_cells(self, table):
        unsure_table = any(
            [True for row in table.rows for cell in row.cells if "?" in cell.text]
        )
        n_rows = len(table.rows)
        n_cols = len(table.columns)
        for row_idx in range(n_rows):
            for col_idx in range(n_cols):
                cell = table.cell(row_idx, col_idx)
                shading_elm = OxmlElement("w:shd")
                if "?" in str(cell.text.strip()):
                    shading_elm.set(qn("w:fill"), "FFFF00")  # yellow color
                elif unsure_table:
                    shading_elm.set(qn("w:fill"), "D3D3D3")  # lightGray color
                cell._tc.get_or_add_tcPr().append(shading_elm)

    def set_column_widths(self, table, table_columns_width):
        for column, width in zip(table.columns, table_columns_width):
            for cell in column.cells:
                try:
                    cell.width = width
                except:
                    pass
